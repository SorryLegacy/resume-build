#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор резюме в формате DOCX
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json
import sys
import subprocess
import os
import shutil


def set_run_font(run, font_name='Calibri', font_size=11, bold=False, italic=False, color=None):
    """Устанавливает шрифт для текста"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Для кириллицы
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_heading(document, text, level=1, size=16, bold=True):
    """Добавляет заголовок"""
    heading = document.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        set_run_font(run, font_size=size, bold=bold)
    return heading


def add_paragraph_with_style(document, text, size=11, bold=False, color=None):
    """Добавляет параграф с заданным стилем"""
    para = document.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, font_size=size, bold=bold, color=color)
    return para


def add_section_header(document, text):
    """Добавляет заголовок раздела"""
    para = document.add_paragraph()
    run = para.add_run(text)
    set_run_font(run, font_size=14, bold=True, color=(0, 51, 102))
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    return para


def add_bullet_point(document, text, indent=0):
    """Добавляет маркированный пункт"""
    para = document.add_paragraph(text, style='List Bullet')
    para.paragraph_format.left_indent = Inches(indent * 0.25)
    for run in para.runs:
        set_run_font(run, font_size=11)
    return para


def generate_resume(data_file='resume_data.json', output_file='resume.docx'):
    """Генерирует резюме из JSON файла"""
    
    # Загружаем данные
    try:
        with open(data_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"Ошибка: файл {data_file} не найден!")
        print("Создайте файл resume_data.json с данными резюме.")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Ошибка при чтении JSON: {e}")
        sys.exit(1)
    
    # Создаем документ
    doc = Document()
    
    # Настройка стилей документа
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Заголовок с именем
    name = data.get('personal_info', {}).get('name', 'Имя Фамилия')
    add_heading(doc, name, level=1, size=20, bold=True)
    
    # Контактная информация
    personal = data.get('personal_info', {})
    contact_info = []
    if personal.get('email'):
        contact_info.append(f"Email: {personal['email']}")
    if personal.get('phone'):
        contact_info.append(f"Телефон: {personal['phone']}")
    if personal.get('location'):
        contact_info.append(f"Местоположение: {personal['location']}")
    if personal.get('linkedin'):
        contact_info.append(f"LinkedIn: {personal['linkedin']}")
    if personal.get('github'):
        contact_info.append(f"GitHub: {personal['github']}")
    
    if contact_info:
        para = doc.add_paragraph(' | '.join(contact_info))
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            set_run_font(run, font_size=10, color=(100, 100, 100))
    
    doc.add_paragraph()  # Пустая строка
    
    # О себе / Профессиональное резюме
    if data.get('summary'):
        add_section_header(doc, 'О себе')
        add_paragraph_with_style(doc, data['summary'], size=11)
        doc.add_paragraph()
    
    # Опыт работы
    if data.get('experience'):
        add_section_header(doc, 'Опыт работы')
        for exp in data['experience']:
            # Название должности и компания
            title_company = f"{exp.get('position', 'Должность')} | {exp.get('company', 'Компания')}"
            para = doc.add_paragraph()
            run = para.add_run(title_company)
            set_run_font(run, font_size=12, bold=True)
            
            # Период работы
            period = exp.get('period', '')
            if period:
                para = doc.add_paragraph()
                run = para.add_run(period)
                set_run_font(run, font_size=10, color=(100, 100, 100), italic=True)
            
            # Описание обязанностей
            if exp.get('description'):
                if isinstance(exp['description'], list):
                    for item in exp['description']:
                        add_bullet_point(doc, item, indent=1)
                else:
                    add_paragraph_with_style(doc, exp['description'], size=11)
            
            doc.add_paragraph()  # Пустая строка между опытом
    
    # Образование
    if data.get('education'):
        add_section_header(doc, 'Образование')
        for edu in data['education']:
            # Учебное заведение и специальность
            degree_school = f"{edu.get('degree', 'Степень')} | {edu.get('school', 'Учебное заведение')}"
            para = doc.add_paragraph()
            run = para.add_run(degree_school)
            set_run_font(run, font_size=12, bold=True)
            
            # Период обучения
            period = edu.get('period', '')
            if period:
                para = doc.add_paragraph()
                run = para.add_run(period)
                set_run_font(run, font_size=10, color=(100, 100, 100), italic=True)
            
            doc.add_paragraph()
    
    # Навыки
    if data.get('skills'):
        add_section_header(doc, 'Навыки')
        skills_text = ', '.join(data['skills']) if isinstance(data['skills'], list) else data['skills']
        add_paragraph_with_style(doc, skills_text, size=11)
        doc.add_paragraph()
    
    # Дополнительные разделы
    if data.get('additional_sections'):
        for section in data['additional_sections']:
            add_section_header(doc, section.get('title', 'Раздел'))
            if isinstance(section.get('content'), list):
                for item in section['content']:
                    if isinstance(item, dict):
                        # Если это словарь, форматируем как ключ-значение
                        key = item.get('key', '')
                        value = item.get('value', '')
                        para = doc.add_paragraph()
                        run1 = para.add_run(f"{key}: ")
                        set_run_font(run1, font_size=11, bold=True)
                        run2 = para.add_run(value)
                        set_run_font(run2, font_size=11)
                    else:
                        add_bullet_point(doc, str(item), indent=1)
            else:
                add_paragraph_with_style(doc, str(section.get('content', '')), size=11)
            doc.add_paragraph()
    
    # Сохраняем документ
    doc.save(output_file)
    print(f"Резюме успешно создано: {output_file}")


def convert_docx_to_pdf_docker(docx_file, pdf_file=None):
    """Конвертирует DOCX в PDF используя LibreOffice в Docker"""
    if pdf_file is None:
        pdf_file = docx_file.replace('.docx', '.pdf')
    
    # Получаем абсолютные пути
    docx_abs = os.path.abspath(docx_file)
    pdf_abs = os.path.abspath(pdf_file)
    work_dir = os.path.dirname(docx_abs)
    docx_name = os.path.basename(docx_abs)
    
    # Проверяем наличие Docker
    if not shutil.which('docker'):
        raise RuntimeError("Docker не установлен или не доступен в PATH")
    
    # Используем образ LibreOffice (можно переопределить через переменную окружения)
    docker_image = os.environ.get('LIBREOFFICE_DOCKER_IMAGE', 'linuxserver/libreoffice:latest')
    
    # Монтируем директорию с файлами и конвертируем
    docker_cmd = [
        'docker', 'run', '--rm',
        '-v', f'{work_dir}:/convert',
        docker_image,
        'libreoffice', '--headless', '--convert-to', 'pdf',
        '--outdir', '/convert',
        f'/convert/{docx_name}'
    ]
    
    try:
        subprocess.run(
            docker_cmd,
            capture_output=True,
            text=True,
            check=True
        )
        
        # LibreOffice создает PDF с тем же именем, но расширением .pdf
        # Переименовываем если нужно
        generated_pdf = os.path.join(work_dir, docx_name.replace('.docx', '.pdf'))
        if os.path.exists(generated_pdf) and generated_pdf != pdf_abs:
            if os.path.exists(pdf_abs):
                os.remove(pdf_abs)
            os.rename(generated_pdf, pdf_abs)
        
        return pdf_abs
    except subprocess.CalledProcessError as e:
        error_msg = e.stderr if e.stderr else e.stdout if e.stdout else str(e)
        raise RuntimeError(f"Ошибка при конвертации через Docker: {error_msg}")
    except FileNotFoundError:
        raise RuntimeError("Docker не найден. Убедитесь, что Docker установлен и доступен.")


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='Генератор резюме в формате DOCX')
    parser.add_argument('-i', '--input', default='resume_data.json',
                        help='Входной JSON файл с данными (по умолчанию: resume_data.json)')
    parser.add_argument('-o', '--output', default='resume.docx',
                        help='Выходной DOCX файл (по умолчанию: resume.docx). Если расширение не указано, автоматически добавляется .docx')
    parser.add_argument('--pdf', action='store_true',
                        help='Также создать PDF версию резюме (требуется Docker с LibreOffice или docx2pdf)')
    
    args = parser.parse_args()
    
    # Автоматически добавляем расширение .docx если его нет
    output_file = args.output
    if not output_file.endswith('.docx'):
        output_file = output_file + '.docx'
    
    generate_resume(args.input, output_file)
    
    # Конвертация в PDF если указан флаг
    if args.pdf:
        pdf_file = output_file.replace('.docx', '.pdf')
        try:
            # Пробуем сначала через Docker (предпочтительный метод)
            convert_docx_to_pdf_docker(output_file, pdf_file)
            print(f"PDF версия создана: {pdf_file}")
        except RuntimeError as e:
            # Если Docker не работает, пробуем через docx2pdf (fallback)
            print(f"Попытка через Docker не удалась: {e}")
            print("Пробую альтернативный метод через docx2pdf...")
            try:
                from docx2pdf import convert
                convert(output_file, pdf_file)
                print(f"PDF версия создана: {pdf_file}")
            except ImportError:
                print("Ошибка: для создания PDF необходимо:")
                print("  1. Установить Docker и образ: docker pull linuxserver/libreoffice:latest")
                print("  2. Или установить библиотеку docx2pdf: pip install docx2pdf")
                sys.exit(1)
            except Exception as e2:
                print(f"Ошибка при конвертации в PDF: {e2}")
                print("\nРекомендации:")
                print("  1. Установите Docker: https://docs.docker.com/get-docker/")
                print("  2. Загрузите образ: docker pull linuxserver/libreoffice:latest")
                print("  3. Или установите LibreOffice локально и используйте docx2pdf")
                sys.exit(1)

